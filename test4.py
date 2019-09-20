from docx import Document
import os
import re
import datetime
os.chdir("C:\\Users\\vprakas\\Desktop\\python\\kpi")
a=os.listdir("C:\\Users\\vprakas\\Desktop\\python\\kpi")
try:
    for i in a:
        cc=i.split("_")[0]
        print(cc) #.................................................
        doc=Document(i)
        chdate=str(doc.tables[0].cell(10,2).text)
        #print(chdate.splitlines())

        for i in chdate.splitlines():
            a=["cst","ct","cdt","central"]
            for x in a:
                if x in i.casefold():
                    changedate=i
                    print(changedate) #................................................
        chdescription=str(doc.tables[0].cell(1,1).text)
        print(chdescription)#..............................................................
        devicename=str(doc.tables[0].cell(9,0).text).replace("Affected Devices:","").replace(" ","")
        print(devicename)#.........................................................
        print("*"*100)
            #if "cst" in i.casefold() or "ct" in i.casefold() or "cdt" in:
            #    print(i)



except:
    pass


