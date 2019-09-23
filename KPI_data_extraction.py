from docx import Document
import os
from dateutil import parser
import pandas as pd
a=os.listdir("C:\\Users\\vprakas\\Desktop\\python\\kpi") #change the directory path to the path that contains all the CRS
os.chdir("C:\\Users\\vprakas\\Desktop\\python\\kpi") #change the directory path to the path that contains all the CRS
lst = []
for i in a:
    try:
        cc=i.split("_")[0] #extracts the cc number
        doc=Document(i)
        chdescription = str(doc.tables[0].cell(1, 1).text) #extracts the change description
        devicename = str(doc.tables[0].cell(9, 0).text) #extracts the affected device name
        if "Affected Devices" not in devicename:
            devicename = str(doc.tables[0].cell(10, 0).text) #extracts the affected device name
        if "Maintenance Window" in str(doc.tables[0].cell(10, 1).text):
            chdate = str(doc.tables[0].cell(10, 2).text)  # extracts the change date in CST
        else:
            chdate = str(doc.tables[0].cell(11, 2).text)  # extracts the change date in CST
        for i in chdate.splitlines():
            a=["cst","ct","cdt","central"]
            for x in a:
                if x in i.casefold():
                    changedate=i
                    try:
                        dt = parser.parse(changedate)
                        chdt=dt.date()
                        time=dt.time()
                    except:
                        chdt=changedate
                        time=changedate
                        pass
        lst.append((cc, chdt, time, chdescription, devicename.split(":")[1]))
    except:
        lst.append(("doc format", "doc format", "doc format", "doc format", "doc format"))
        print("doc format",i)
        pass
df=pd.DataFrame(lst)
df.to_excel("C:\\Users\\vprakas\\Desktop\\python\\KPI_output.xlsx") #Give any path where you require the output excel sheet to be generated