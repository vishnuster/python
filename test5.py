from docx import Document
import os
from dateutil import parser
import pandas as pd
os.chdir("C:\\Users\\vprakas\\Desktop\\python\\kpi")
a=os.listdir("C:\\Users\\vprakas\\Desktop\\python\\kpi")
lst = []
for i in a:
    try:
        cc=i.split("_")[0]
        doc=Document(i)
        chdescription = str(doc.tables[0].cell(1, 1).text)
        devicename = str(doc.tables[0].cell(9, 0).text).replace("Affected Devices:", "").replace(" ", "")
        chdate=str(doc.tables[0].cell(10,2).text)
        for i in chdate.splitlines():
            a=["cst","ct","cdt","central"]
            for x in a:
                if x in i.casefold():
                    changedate=i
                    ###############################################
                    try:
                        dt = parser.parse(changedate)
                        chdt=dt.date()
                        time=dt.time()
                    except:
                        chdt=changedate
                        time=changedate
                        pass
        lst.append((cc, chdt, time, chdescription, devicename))
    except:
        er="error"
        #lst.append((cc, chdt, time, chdescription, devicename))
        lst.append(("doc format", "doc format", "doc format", "doc format", "doc format"))
        print("doc format",i)
        pass
#print(lst)
df=pd.DataFrame(lst)
#print(df)
df.to_excel("C:\\Users\\vprakas\\Desktop\\python\\KPI_output.xlsx")